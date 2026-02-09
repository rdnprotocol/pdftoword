"""
Export module for converting corrected text to Word document.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def export_to_word(text: str, output_path: str) -> str:
    """
    Export text to Word document (.docx).
    
    Args:
        text: Text content to export
        output_path: Output file path (should end with .docx)
        
    Returns:
        Path to created file
    """
    # Create document
    doc = Document()
    
    # Set default font for Mongolian Cyrillic
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add title
    title = doc.add_heading('Зассан Бичвэр', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add text content
    paragraphs = text.split('\n\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            para.style = doc.styles['Normal']
    
    # Save document
    if not output_path.endswith('.docx'):
        output_path += '.docx'
    
    doc.save(output_path)
    
    return output_path
